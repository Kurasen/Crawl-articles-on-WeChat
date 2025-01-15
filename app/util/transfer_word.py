import json
import os
from docx import Document
from docx.shared import Inches

# 读取 JSON 文件
input_json_path = r'D:\Technology\Code\Python\Crawl articles on WeChat\data\word\土壤观察\fixed_articles.json'
output_dir = r'D:\Technology\Code\Python\Crawl articles on WeChat\data\output'
image_base_path = r'D:\Technology\Code\Python\Crawl articles on WeChat\data\images\土壤观察'

# 不允许的字符列表
invalid_characters = ['|', '<', '>', ':', '"', '/', '\\', '?', '*']

# 替换文件名中的非法字符
def sanitize_filename(filename):
    for char in invalid_characters:
        filename = filename.replace(char, '_')
    return filename

def fix_image_path(image_path):
    # 如果路径以 'data\\images\\土壤观察' 开头，移除这一部分
    prefix = 'data\\images\\土壤观察'
    if image_path.startswith(prefix):
        return image_path[len(prefix):].lstrip("\\/")
    return image_path

with open(input_json_path, 'r', encoding='utf-8') as file:
    articles = json.load(file)

# 输出目录不存在则创建
if not os.path.exists(output_dir):
    os.makedirs(output_dir)

# 转换每篇文章
for article in articles:
    doc = Document()
    doc.add_heading(article['title'], 0)
    doc.add_paragraph(f"链接：{article['link']}")
    doc.add_paragraph(f"创建时间：{article['create_time']}")
    doc.add_paragraph()  # 空行

    # 处理文章内容
    for item in article['content']:
        if item['type'] == 'text':
            doc.add_paragraph(item['content'])
        elif item['type'] == 'images':
            image_path = item['content']
            # 修正图片路径
            image_full_path = os.path.join(image_base_path, fix_image_path(image_path))
            #print(f"原始图片路径：{image_path}")
            #print(f"修正后完整路径：{image_full_path}")
            if os.path.exists(image_full_path):
                doc.add_picture(image_full_path, width=Inches(4))  # 调整图片大小
                doc.add_paragraph()  # 空行
            else:
                print(f"图片 {image_full_path} 未找到，跳过。")

    # 生成合法的文件名
    sanitized_title = sanitize_filename(article['title'])
    output_doc_path = os.path.join(output_dir, f"{sanitized_title}.docx")
    doc.save(output_doc_path)
    #print(f"已保存文章 '{article['title']}' 到 {output_doc_path}")

print("所有文章已成功转换为 Word 文档。")
